
# --------------------------------------------------------------------------------------------------
# --------------------------------------------------------------------------------------------------
# --------------------------------------- Docx Testing ---------------------------------------------
# --------------------------------------------------------------------------------------------------
# --------------------------------------------------------------------------------------------------

import docx
import json

### User Input ------------------------------------------------------ ###
# Identifying the file location of the template
file_target_name = 'Target Resume Template2.docx'  # with extension
file = 'target_resumes/' + file_target_name

# Identifying an output file location (moved for demo)
# fileout = 'output_resumes/' + 'Output Testing.docx'

### Target Resume Completeion -------------------------------- ###
# Import InfoDict json file to add to target resume template
with open('output_resumes/infoDict.txt') as json_file:
    infoDict = json.load(json_file)

new = infoDict

new['#NAME'] = infoDict['name']['first'] + " " + infoDict['name']['middle']\
    + " " + infoDict['name']['last']
new['#LASTNAME'] = infoDict['name']['last']
new['#EDUCATION'] = infoDict['peducation']['degree'] + ',' + infoDict['peducation']['field'] + \
    ',' + infoDict['peducation']['college'] + \
    ',' + infoDict['peducation']['year']

# converts the output list into a dictionary to work with the below
testing = {}
for x in range(len(infoDict['gtExp']['engagements'])):
    key = 'EXPERIENCE' + str(x)
    testing[key] = infoDict['gtExp']['engagements'][x]

# Identifying an output file location (for demo)
fileout = 'output_resumes/' + \
    infoDict['name']['first'] + ' ' + \
    infoDict['name']['last'] + ' ' + file_target_name + '.docx'

# --------------------------------------------------------------------------------------------------
# --------------------------------- Duplicates tags for repetitive fields --------------------------
# --------------------------------------------------------------------------------------------------

# Read in Target resume template
doc = docx.Document(file)  # Does not work with tables yet

# returns the numbers for paragraphs that contain an experience related field


def getParaNumbers(file):
    para = []
    # Loop through fields in dictionary
    for field in testing['EXPERIENCE0']:
        count = -1
        # Loop through paragraphs in document
        for p in doc.paragraphs:
            count = count + 1
            # If field is in paragraph, loop through the runs
            if field in p.text:
                para.append(count)
    return para


para = getParaNumbers(file)

# removing duplicates
para = list(dict.fromkeys(para))
para.sort()

# Gets formatting properties of template runs - will need to add more


def clone_run_props(tmpl_run, this_run):
    this_run.bold = tmpl_run.bold
    this_run.italic = tmpl_run.italic
    this_run.underline = tmpl_run.underline
    this_run.font.name = tmpl_run.font.name
    this_run.font.size = tmpl_run.font.size
    this_run.font.color.rgb = tmpl_run.font.color.rgb


# only creates new paragraphs at bottom of document
# For each experience in the dictionary
for x in range(0, len(testing)-1):
    # Add a blank paragraph
    doc.add_paragraph()
    # For each paragraph in the copy range
    for i in para:
        # point to the copy paragraph
        template_paragraph = doc.paragraphs[i]
        # add a new paragraph
        new_paragraph = doc.add_paragraph()
        #new_paragraph.paragraph_format = template_paragraph.paragraph_format
        count = 0
        # for each run in the copy paragraph
        for run in template_paragraph.runs:
            # add a new run in the new paragaph
            cloned_run = new_paragraph.add_run()
            # assign the properties of the copy to the new
            clone_run_props(run, cloned_run)
            # assigns the text of the run
            cloned_run.text = doc.paragraphs[i].runs[count].text
            count = count + 1

# --------------------------------------------------------------------------------------------------
# --------------------------------------- replace tags ---------------------------------------------
# --------------------------------------------------------------------------------------------------

# Replace strings for non repetitive fields


def replace_string(file):
    # Open file
    #doc = docx.Document(file)
    # Loop through fields in dictionary
    for field in new:
        if field != 'gtExp':
            # Loop through paragraphs in document
            for p in doc.paragraphs:
                # If field is in paragraph, loop through the runs
                if field in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                        # Replace the key with the desired text
                        if field in inline[i].text:
                            text = inline[i].text.replace(field, new[field])
                            inline[i].text = text
            # Loop through the tables in the document
            for table in doc.tables:
                # Loop through the rows in the table
                for row in table.rows:
                    # Loop through the cells in the row
                    for cell in row.cells:
                        # Loop through the paragraphs in the cell
                        for p in cell.paragraphs:
                            # If feild is in paragraph, loop through the runs
                            if field in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    # Replace the key with the desired text
                                    if field in inline[i].text:
                                        text = inline[i].text.replace(
                                            field, new[field])
                                        inline[i].text = text

    # Replace the now duplicated tags with the appropriate values
    # Loop through keys in dictionary
    for key in testing:
        # Loop through values in each key
        for field in testing[key]:
                # Loop through paragraphs in document
            for p in doc.paragraphs:
                        # If field is in paragraph, loop through the runs
                if field in p.text:
                    inline = p.runs
                    for i in range(len(inline)):
                                # Replace the key with the desired text
                        if field in inline[i].text:
                            text = inline[i].text.replace(
                                field, testing[key][field])
                            inline[i].text = text
                            break
                    break
         # Loop through the tables in the document
            for table in doc.tables:
                    # Loop through the rows in the table
                for row in table.rows:
                    # Loop through the cells in the row
                    for cell in row.cells:
                        # Loop through the paragraphs in the cell
                        for p in cell.paragraphs:
                            # If feild is in paragraph, loop through the runs
                            if field in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                        # Replace the key with the desired text
                                    if field in inline[i].text:
                                        text = inline[i].text.replace(
                                            field, testing[key][field])
                                        inline[i].text = text
                                        break

    # Save to fileout
    doc.save(fileout)
    return


replace_string(file)

# --------------------------------------------------------------------------------------------------
# --------------------------------------------------------------------------------------------------
# ------------------------------------------ End Script --------------------------------------------
# --------------------------------------------------------------------------------------------------
# --------------------------------------------------------------------------------------------------

# --------------------------------------------------------------------------------------------------
# ---------------------------------------- Merging two files ---------------------------------------
# --------------------------------------------------------------------------------------------------
'''
# combines files while maintaining formatting
## Needs to be adjusted to use all files that were created. Start each resume on new page?
doc1 = docx.Document('Output Testing.docx')
doc2 = docx.Document('Output Testing2.docx')
for element in doc2.element.body:
    doc1 .element.body.append(element)
doc1.save('merge.docx')
'''

'''
# Example of data.. would be replaced by the output from the scraping
new = {'#NAME':'Brian C. Sullivan',
       '#EDUCATION':'B.B.A., Economics, James Madison University, 2018, Minor / Concentration: Concentration in Financial Economics',
       '#TRAININGS':'Microsoft Office Specialist Excel 2013, Bloomberg Market Concept',
       '#CLEARANCE':'',
       '#YEARSEXPERIENCE':'> 1 Year',
       '#INTRODUCTION': 'Mr. Sullivan is an Advisory Associate with Grant Thornton LLP’s Public Sector Practice in Arlington, VA and is aligned with the Advanced Digital Technology and Analytics service line within the practice. He holds a Bachelor of Business Administration (B.B.A.) degree in Economics with a concentration in Financial Economics from James Madison University. During his time as Business Analyst, he has gained experience in data analysis and visualization and supporting evaluations of financial management systems and processes. Through his education and work experience, Mr. Sullivan has developed skills in data visualization, data analytics, econometric modeling, program evaluation, and client services. He has experience with computer programs such as Microsoft PowerBI, Tableau, Qlik Sense, SAS, and the MS Office Suite, as well as open-source programming languages such as Python, R, SQL, VBA and HTML.',
       #'#EXPERIENCE': 'Mr. Sullivan is a member of the Tableau Visualization Support team within OPIA. This effort includes utilizing Tableau, a data visualization tool, to create various reports required both internally for OPIA and USPTO stakeholders at various leadership levels as well as for the external consumption by the public. Mr. Sullivan leads the requirements gathering, development and refresh processes of all reports and more than 12 dashboards from the data validation and staging process to the final deliverable. Mr. Sullivan also developed numerous VBA Macros and Python scripts to include ETL scripts, semi-automated reporting tools, and a PDF scraping tool in order to support dashboard refresh processes, data collection efforts and ad hoc reporting requests from OPIA stakeholders.',
       #'#AGENCY':'Department of Commerce, United States Patent and Trademark Office (USPTO), Office of Policy and International Affairs (OPIA)',
       #'#PROJECTDATES': '(February 2019 – Present)',
       '#TITLES':'',
       '#SUMMARY':'Experienced data analytics professional with demonstrated ability developing Key Performance Indicators (KPI), data mining, data preparation, and data visualization.',
       '#CERTIFICATIONS':'',
       '#LASTNAME':'Sullivan',
       '#TITLE':''}

testing = {'EXPERIENCE1':{
               '#AGENCY':'Department of Commerce, United States Patent and Trademark Office (USPTO), Office of Policy and International Affairs (OPIA)',
               '#PROJECTDATES': '(February 2019 – Present)',
               '#EXPERIENCE': 'Mr. Sullivan is a member of the Tableau Visualization Support team within OPIA. This effort includes utilizing Tableau, a data visualization tool, to create various reports required both internally for OPIA and USPTO stakeholders at various leadership levels as well as for the external consumption by the public. Mr. Sullivan leads the requirements gathering, development and refresh processes of all reports and more than 12 dashboards from the data validation and staging process to the final deliverable. Mr. Sullivan also developed numerous VBA Macros and Python scripts to include ETL scripts, semi-automated reporting tools, and a PDF scraping tool in order to support dashboard refresh processes, data collection efforts and ad hoc reporting requests from OPIA stakeholders.'},
           'EXPERIENCE2':{
               '#AGENCY':'State of Texas Department of Transportation (TxDOT)',
               '#PROJECTDATES':'(May 2019 – August 2019)',
               '#EXPERIENCE': 'Grant Thornton was contracted to develop a study in support of TxDOT on the ten-year Enterprise Information Management (EIM) Strategic Plan for all statewide systems and processes, including programs that address data governance, system architecture, advanced analytics, employee training, and effective communications. As part of this effort, Mr. Sullivan led the Quality Control process in support of over 350 project artifacts across more than 60 projects.'},
            'EXPERIENCE3':{
               '#AGENCY':'Library of Congress, United States Copyright Office (USCO), Licensing Division',
               '#PROJECTDATES':'(August 2018 – January 2019)',
               '#EXPERIENCE': 'Mr. Sullivan supported the Library of Congress United States Copyright Office (USCO) Accounting System Analysis of Alternatives engagement. The project provided an analysis based recommendation for alternatives to the legacy accounting and investment management systems. In this role Mr. Sullivan supported the client’s needs by conducting stakeholder interviews, conducting market research, developing process charts and evaluating alternatives. Mr. Sullivan also supported the creation of deliverables that support the continuing communication between the client and GT. The team was able to provide an initial list of COTS alternatives to the current process, and identify weaknesses in the current processes and systems. Mr. Sullivan assisted in the evaluations of cost and capability of potential solutions, which fed the recommendations.'}
        }
'''
